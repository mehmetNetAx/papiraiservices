from docx import Document

def fill_document01(template_path, output_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key,value)
                for run in paragraph.runs:
                    run.text = run.text.replace(key,value)
    
    doc.save(output_path)



def replace_text_in_paragraphs(paragraphs, data):
    for paragraph in paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key,value)
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

def replace_text_in_tables(tables, data):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, data)

def fill_document(template_path, output_path, data):
    doc = Document(template_path)
    
    # Replace keys in paragraphs
    replace_text_in_paragraphs(doc.paragraphs, data)
    
    # Replace keys in tables
    replace_text_in_tables(doc.tables, data)

    
    doc.save(output_path)




data = {
        '<tarih>' : '20.06.2024',
        '<Sirket Unvani>' : 'NETAX BİLİŞİM HİZMETLERİ',
        '<Sirket Adresi>' : 'Altunizade Mah. Üsküdar/İstanbul',
        '<ticaretsicilno>' : '45730-5',
        '<firmawebadresi>' : 'www.netaxtech.com',
        '<firmaKEPadresi>' :  '22222',
        '<sozlesmenin_konusu>' : 'Yazilim Geliştirme Hizmeti',
        '<sozlesme_toplam_bedeli>' : '3000.000.- TRL',
        '<sözleme_bedeli_maddesi_1>' : '%30 Peşin geri kalan eşit taksitle',
        '<firma_banka_adi>' : 'VakifBank',
        '<firma_banka_sube_adi>' : 'Kasimpasa Subesi',
        '<firma_hesap_no>' : '222222 3434',
        '<firma_iban>' : 'TR20 1111 2222 3333 4444 55',
        '<cezai_sart_1>' : 'Cezai Şart 1',
        '<cezai_sart_2>' : 'Cezai Şart 2',
        '<cezai_sart_3>' : 'Cezai Şart 3',
        
        '<esgis_istenmeyen_maddeler_1>' : 'İstenmeyen 1',
        '<esgis_istenmeyen_maddeler_2>' : 'İstenmeyen 2',
        '<esgis_istenmeyen_maddeler_3>' : 'İstenmeyen 3', 
        '<esgis_istenmeyen_maddeler_4>' : 'İstenmeyen 4',

        '<kac_nüsha>' : '2 Nüsha',
        '<sozlesme_suresi>' : '1 Yıl',
        '<firma_ticaret_unvani>' : 'Netax Bilişim ve Danışmanlık Hizmetleri Tic. Ltd. Sti.',
        '<FOimza1>' : 'imza Yönetici 1',
        '<FOimza2>' : 'imza Yönetici 2',
        '<FOimza3>' : 'imza Yönetici 3',
        '<firmaimza1>' : 'Firma imza 1',
        '<firmaimza2>' : 'Firma imza 2',
        '<firmaimza3>' : 'Firma imza 3'

}

template_path = 'FO_Contract_Template.docx'
output_path = 'FO_Filled_Contract.docx'

fill_document(template_path, output_path, data)

